# -*- coding: utf-8 -*-
from docx import Document
import sys

doc = Document('D:/相关资料/openclaw/（整合版）PRD_监督与调整模块.docx')

# 输出段落
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(p.text)

# 输出表格
for t in doc.tables:
    print("\n--- TABLE ---")
    for row in t.rows:
        print(" | ".join([cell.text for cell in row.cells]))
